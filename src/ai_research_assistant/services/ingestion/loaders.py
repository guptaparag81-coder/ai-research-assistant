"""Text extraction and metadata extraction for supported document types."""

import io
from dataclasses import dataclass

import docx
from pypdf import PdfReader

from ai_research_assistant.core.exceptions import IngestionError, UnsupportedFileTypeError
from ai_research_assistant.db.models.document import DocumentType

_CONTENT_TYPE_TO_DOCUMENT_TYPE: dict[str, DocumentType] = {
    "application/pdf": DocumentType.PDF,
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DocumentType.DOCX,
    "text/plain": DocumentType.TXT,
    "text/markdown": DocumentType.MARKDOWN,
}


@dataclass(frozen=True, slots=True)
class ExtractedDocument:
    text: str
    title: str | None
    page_count: int | None
    document_type: DocumentType


def resolve_document_type(content_type: str) -> DocumentType:
    document_type = _CONTENT_TYPE_TO_DOCUMENT_TYPE.get(content_type)
    if document_type is None:
        raise UnsupportedFileTypeError(f"Unsupported content type: {content_type}")
    return document_type


def _extract_pdf(raw: bytes) -> ExtractedDocument:
    reader = PdfReader(io.BytesIO(raw))
    pages = [page.extract_text() or "" for page in reader.pages]
    text = "\n\n".join(pages).strip()
    title = None
    if reader.metadata is not None:
        title = reader.metadata.title
    return ExtractedDocument(
        text=text, title=title, page_count=len(reader.pages), document_type=DocumentType.PDF
    )


def _extract_docx(raw: bytes) -> ExtractedDocument:
    document = docx.Document(io.BytesIO(raw))
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    text = "\n\n".join(paragraphs).strip()
    core_props = document.core_properties
    title = core_props.title or None
    return ExtractedDocument(
        text=text, title=title, page_count=None, document_type=DocumentType.DOCX
    )


def _extract_plain_text(raw: bytes, document_type: DocumentType) -> ExtractedDocument:
    text = raw.decode("utf-8", errors="replace").strip()
    return ExtractedDocument(text=text, title=None, page_count=None, document_type=document_type)


def extract_document(*, raw: bytes, content_type: str) -> ExtractedDocument:
    """Extract normalized text and metadata from raw file bytes."""
    document_type = resolve_document_type(content_type)
    try:
        if document_type is DocumentType.PDF:
            extracted = _extract_pdf(raw)
        elif document_type is DocumentType.DOCX:
            extracted = _extract_docx(raw)
        else:
            extracted = _extract_plain_text(raw, document_type)
    except Exception as exc:
        raise IngestionError(f"Failed to parse {document_type.value} document") from exc

    if not extracted.text:
        raise IngestionError("Document contains no extractable text")
    return extracted
